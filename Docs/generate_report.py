"""
Script tạo báo cáo Word giới thiệu phần mềm Quản lý VKTBKT
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_report():
    doc = Document()
    
    # ===== THIẾT LẬP FONT MẶC ĐỊNH =====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # ===== TRANG BÌA =====
    # Tiêu đề phụ
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BÁO CÁO SẢN PHẨM PHẦN MỀM")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tiêu đề chính
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PHẦN MỀM QUẢN LÝ VŨ KHÍ TRANG BỊ KỸ THUẬT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("THÔNG QUA CÔNG NGHỆ MÃ QR")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("(QuanLyVKTBKT)")
    run.italic = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Thông tin
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phiên bản: 1.0.0")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Năm 2024")
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ===== MỤC LỤC =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MỤC LỤC")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    toc_items = [
        ("1. GIỚI THIỆU TỔNG QUAN", "3"),
        ("   1.1. Bối cảnh và nhu cầu", "3"),
        ("   1.2. Mục tiêu của phần mềm", "3"),
        ("   1.3. Phạm vi ứng dụng", "4"),
        ("2. CÔNG NGHỆ SỬ DỤNG", "5"),
        ("   2.1. Ngôn ngữ lập trình và Framework", "5"),
        ("   2.2. Cơ sở dữ liệu", "5"),
        ("   2.3. Thư viện xử lý QR Code", "6"),
        ("   2.4. Thư viện xuất báo cáo", "6"),
        ("3. KIẾN TRÚC HỆ THỐNG", "7"),
        ("   3.1. Mô hình MVC", "7"),
        ("   3.2. Cấu trúc thư mục", "7"),
        ("   3.3. Sơ đồ cơ sở dữ liệu", "8"),
        ("4. CÁC CHỨC NĂNG CHÍNH", "9"),
        ("   4.1. Quản lý thiết bị", "9"),
        ("   4.2. Quản lý mã QR", "10"),
        ("   4.3. Quản lý bảo dưỡng", "11"),
        ("   4.4. Quản lý cho mượn", "11"),
        ("   4.5. Quản lý đơn vị", "12"),
        ("   4.6. Quản lý người dùng", "12"),
        ("   4.7. Xuất báo cáo PDF", "13"),
        ("5. ĐIỂM NỔI BẬT CỦA HỆ THỐNG", "14"),
        ("6. HƯỚNG DẪN CÀI ĐẶT VÀ SỬ DỤNG", "16"),
        ("7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "18"),
    ]
    
    for item, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(item).font.name = 'Times New Roman'
        p.add_run(" " + "." * (60 - len(item)) + " ").font.name = 'Times New Roman'
        p.add_run(page).font.name = 'Times New Roman'
    
    doc.add_page_break()
    
    # ===== PHẦN 1: GIỚI THIỆU TỔNG QUAN =====
    h1 = doc.add_heading('1. GIỚI THIỆU TỔNG QUAN', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('1.1. Bối cảnh và nhu cầu', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("""Trong công tác quản lý vũ khí, khí tài tại các đơn vị quân đội hiện nay vẫn còn nhiều hạn chế:""")
    
    bullets = [
        "Quản lý chủ yếu bằng sổ sách, giấy tờ thủ công, dễ thất lạc và khó tra cứu",
        "Khó khăn trong việc theo dõi tình trạng, lịch sử bảo dưỡng của từng thiết bị",
        "Mất nhiều thời gian khi cần kiểm kê hoặc tìm kiếm thông tin",
        "Thiếu công cụ hỗ trợ báo cáo, thống kê nhanh chóng",
        "Khó quản lý việc cho mượn, điều chuyển thiết bị giữa các đơn vị"
    ]
    
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("""Trước yêu cầu chuyển đổi số và hiện đại hóa công tác quản lý, việc xây dựng một phần mềm quản lý vũ khí trang bị kỹ thuật (VKTBKT) là hết sức cần thiết và cấp bách.""")
    
    h2 = doc.add_heading('1.2. Mục tiêu của phần mềm', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Phần mềm Quản lý VKTBKT được phát triển với các mục tiêu:").bold = True
    
    goals = [
        "Số hóa toàn bộ hồ sơ vũ khí, khí tài: Lưu trữ đầy đủ thông tin về từng thiết bị bao gồm tên, số hiệu, loại, tình trạng, đơn vị quản lý",
        "Ứng dụng công nghệ mã QR: Mỗi thiết bị được gán một mã QR định danh duy nhất, cho phép tra cứu nhanh chóng bằng camera",
        "Theo dõi lịch sử bảo dưỡng: Ghi lại đầy đủ các hoạt động sửa chữa, bảo dưỡng, giúp lập kế hoạch bảo trì hiệu quả",
        "Quản lý cho mượn/điều chuyển: Theo dõi chính xác thiết bị được mượn bởi đơn vị nào, thời gian mượn/trả",
        "Hỗ trợ xuất báo cáo: Tự động tạo các báo cáo PDF phục vụ công tác kiểm kê, báo cáo cấp trên",
        "Đảm bảo an toàn dữ liệu: Hoạt động offline, không kết nối internet, dữ liệu được lưu trữ cục bộ"
    ]
    
    for i, goal in enumerate(goals, 1):
        p = doc.add_paragraph(f"{goal}", style='List Bullet')
    
    h2 = doc.add_heading('1.3. Phạm vi ứng dụng', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("""Phần mềm được thiết kế để ứng dụng tại các đơn vị quân đội từ cấp đại đội trở lên, phục vụ cho:""")
    
    scopes = [
        "Cán bộ quân khí: Quản lý hồ sơ, kiểm kê thiết bị",
        "Nhân viên kỹ thuật: Ghi nhật ký bảo dưỡng, sửa chữa",
        "Chỉ huy đơn vị: Tra cứu, xem báo cáo tổng hợp",
        "Quản trị viên: Quản lý tài khoản, phân quyền người dùng"
    ]
    
    for scope in scopes:
        p = doc.add_paragraph(scope, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 2: CÔNG NGHỆ SỬ DỤNG =====
    h1 = doc.add_heading('2. CÔNG NGHỆ SỬ DỤNG', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('2.1. Ngôn ngữ lập trình và Framework', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    # Bảng công nghệ
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Công nghệ', 'Phiên bản', 'Mô tả']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '1976D2')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = None
    
    data = [
        ['Python', '3.10+', 'Ngôn ngữ lập trình chính, mạnh mẽ và dễ phát triển'],
        ['PyQt6', '6.6.0+', 'Framework xây dựng giao diện đồ họa (GUI) hiện đại'],
        ['SQLite', '3.x', 'Hệ quản trị CSDL nhúng, nhẹ và không cần cài đặt server'],
    ]
    
    for i, row_data in enumerate(data, 1):
        cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            cells[j].text = cell_text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Python ")
    run.bold = True
    p.add_run("được lựa chọn vì tính đơn giản, dễ bảo trì và có cộng đồng hỗ trợ lớn. ")
    run = p.add_run("PyQt6 ")
    run.bold = True
    p.add_run("cho phép xây dựng giao diện desktop đẹp, đáp ứng (responsive) và hoạt động đa nền tảng.")
    
    h2 = doc.add_heading('2.2. Cơ sở dữ liệu', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    run = p.add_run("SQLite ")
    run.bold = True
    p.add_run("""là lựa chọn tối ưu cho ứng dụng này vì:""")
    
    sqlite_benefits = [
        "Không cần cài đặt database server riêng",
        "Dữ liệu được lưu trong một file duy nhất, dễ sao lưu và di chuyển",
        "Hiệu năng tốt với lượng dữ liệu vừa và nhỏ (phù hợp cấp đơn vị)",
        "Hoạt động hoàn toàn offline, đảm bảo an toàn thông tin quân sự"
    ]
    
    for benefit in sqlite_benefits:
        p = doc.add_paragraph(benefit, style='List Bullet')
    
    h2 = doc.add_heading('2.3. Thư viện xử lý QR Code', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Thư viện', 'Phiên bản', 'Chức năng']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '4CAF50')
    
    qr_libs = [
        ['qrcode', '7.4.2+', 'Tạo mã QR từ dữ liệu thiết bị'],
        ['pyzbar', '0.1.9+', 'Giải mã QR từ hình ảnh/camera'],
        ['OpenCV', '4.8.0+', 'Xử lý hình ảnh và điều khiển camera'],
    ]
    
    for i, row_data in enumerate(qr_libs, 1):
        cells = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            cells[j].text = cell_text
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('2.4. Thư viện xuất báo cáo', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    run = p.add_run("ReportLab ")
    run.bold = True
    p.add_run("""(phiên bản 4.0+) được sử dụng để tạo các báo cáo PDF chuyên nghiệp với các tính năng:""")
    
    reportlab_features = [
        "Hỗ trợ font tiếng Việt (Arial, Times New Roman)",
        "Tạo bảng biểu với định dạng linh hoạt",
        "Nhúng hình ảnh QR code vào báo cáo",
        "Xuất PDF theo nhiều định dạng (A4, A4 ngang)"
    ]
    
    for feature in reportlab_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== PHẦN 3: KIẾN TRÚC HỆ THỐNG =====
    h1 = doc.add_heading('3. KIẾN TRÚC HỆ THỐNG', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('3.1. Mô hình MVC (Model-View-Controller)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("""Phần mềm được xây dựng theo mô hình MVC, giúp tách biệt các thành phần và dễ dàng bảo trì:""")
    
    mvc_items = [
        ("Model (Mô hình dữ liệu)", "Định nghĩa cấu trúc dữ liệu và tương tác với CSDL. Bao gồm: Equipment, MaintenanceLog, LoanLog, Unit, User, Category"),
        ("View (Giao diện)", "Hiển thị thông tin và nhận tương tác từ người dùng. Sử dụng PyQt6 để xây dựng các cửa sổ, dialog, bảng dữ liệu"),
        ("Controller (Điều khiển)", "Xử lý logic nghiệp vụ, làm cầu nối giữa Model và View. Kiểm tra dữ liệu đầu vào, thực hiện các thao tác CRUD")
    ]
    
    for title, desc in mvc_items:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)
    
    h2 = doc.add_heading('3.2. Cấu trúc thư mục dự án', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    structure = """
QuanLyVKTB/
├── main.py                 # File khởi chạy ứng dụng
├── requirements.txt        # Danh sách thư viện
├── data/                   # Dữ liệu
│   ├── vktbkt.db          # Database SQLite
│   ├── qr_codes/          # Ảnh mã QR
│   └── exports/           # File PDF xuất
├── assets/                 # Tài nguyên (icons, fonts)
└── src/                    # Source code
    ├── config.py          # Cấu hình ứng dụng
    ├── models/            # Data models
    ├── views/             # Giao diện người dùng
    ├── controllers/       # Logic xử lý
    └── services/          # Các dịch vụ (QR, Camera, Export)
"""
    
    p = doc.add_paragraph()
    p.add_run(structure).font.name = 'Consolas'
    p.add_run(structure).font.size = Pt(10)
    
    h2 = doc.add_heading('3.3. Sơ đồ cơ sở dữ liệu', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống sử dụng các bảng dữ liệu chính sau:").bold = True
    
    tables_desc = [
        ("equipment", "Lưu thông tin thiết bị: tên, số hiệu, loại, tình trạng, đơn vị, vị trí"),
        ("units", "Quản lý cây đơn vị: tên, mã, cấp trên, cấp độ, chỉ huy"),
        ("users", "Tài khoản người dùng: username, password (mã hóa), vai trò"),
        ("maintenance_log", "Nhật ký bảo dưỡng: thiết bị, loại công việc, kỹ thuật viên, ngày"),
        ("loan_log", "Lịch sử cho mượn: thiết bị, đơn vị mượn, ngày mượn/trả"),
        ("categories", "Danh mục loại trang bị: súng, khí tài, phương tiện"),
        ("maintenance_types", "Loại công việc bảo dưỡng: kiểm tra, sửa chữa, bảo trì")
    ]
    
    for table_name, desc in tables_desc:
        p = doc.add_paragraph()
        run = p.add_run(f"• {table_name}: ")
        run.bold = True
        run.font.name = 'Consolas'
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ===== PHẦN 4: CÁC CHỨC NĂNG CHÍNH =====
    h1 = doc.add_heading('4. CÁC CHỨC NĂNG CHÍNH', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('4.1. Quản lý thiết bị (Equipment Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Đây là chức năng cốt lõi của hệ thống, cho phép:").bold = True
    
    equip_features = [
        "Thêm mới thiết bị: Nhập đầy đủ thông tin gồm tên, số hiệu (duy nhất), loại, nhà sản xuất, năm sản xuất, đơn vị quản lý, vị trí, mô tả",
        "Xem danh sách: Hiển thị bảng với các cột thông tin chính, hỗ trợ phân trang",
        "Tìm kiếm và lọc: Tìm theo tên, số hiệu; lọc theo loại, tình trạng, đơn vị",
        "Chỉnh sửa thông tin: Cập nhật bất kỳ trường thông tin nào của thiết bị",
        "Xóa thiết bị: Xóa mềm (đánh dấu không hoạt động) để lưu lịch sử",
        "Xem chi tiết: Hiển thị đầy đủ thông tin kèm mã QR và lịch sử"
    ]
    
    for feature in equip_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("Phân loại tình trạng thiết bị theo 5 cấp:").italic = True
    
    status_table = doc.add_table(rows=6, cols=2)
    status_table.style = 'Table Grid'
    
    status_data = [
        ('Cấp độ', 'Ý nghĩa'),
        ('Cấp 1', 'Tốt - Sẵn sàng chiến đấu'),
        ('Cấp 2', 'Khá - Hoạt động bình thường'),
        ('Cấp 3', 'Trung bình - Cần theo dõi'),
        ('Cấp 4', 'Yếu - Cần bảo dưỡng'),
        ('Cấp 5', 'Kém - Cần sửa chữa lớn hoặc thanh lý'),
    ]
    
    for i, (col1, col2) in enumerate(status_data):
        cells = status_table.rows[i].cells
        cells[0].text = col1
        cells[1].text = col2
        if i == 0:
            cells[0].paragraphs[0].runs[0].bold = True
            cells[1].paragraphs[0].runs[0].bold = True
            set_cell_shading(cells[0], 'E3F2FD')
            set_cell_shading(cells[1], 'E3F2FD')
    
    doc.add_paragraph()
    
    h2 = doc.add_heading('4.2. Quản lý mã QR (QR Code Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    qr_features = [
        ("Tự động sinh mã QR", "Khi thêm thiết bị mới, hệ thống tự động tạo mã QR duy nhất chứa thông tin: ID thiết bị, số hiệu"),
        ("Quét mã QR bằng camera", "Sử dụng webcam để quét mã QR, tự động tra cứu và hiển thị thông tin thiết bị"),
        ("Lưu ảnh mã QR", "Xuất ảnh QR định dạng PNG để in ấn"),
        ("In trực tiếp", "Hỗ trợ xem trước và in mã QR ra giấy decal"),
        ("Xuất bảng mã QR", "Tạo file PDF chứa nhiều mã QR để in hàng loạt")
    ]
    
    for title, desc in qr_features:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)
    
    h2 = doc.add_heading('4.3. Quản lý bảo dưỡng (Maintenance Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    maintenance_features = [
        "Ghi nhật ký bảo dưỡng cho từng thiết bị",
        "Phân loại công việc: Kiểm tra định kỳ, Sửa chữa, Bảo dưỡng, Thay thế linh kiện",
        "Ghi nhận kỹ thuật viên thực hiện",
        "Theo dõi trạng thái: Đang thực hiện, Hoàn thành",
        "Cập nhật tình trạng thiết bị sau bảo dưỡng",
        "Xem lịch sử đầy đủ theo thiết bị hoặc theo thời gian"
    ]
    
    for feature in maintenance_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    h2 = doc.add_heading('4.4. Quản lý cho mượn (Loan Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    loan_features = [
        "Tạo phiếu cho mượn thiết bị",
        "Ghi nhận đơn vị mượn, ngày mượn, số điện thoại liên hệ",
        "Theo dõi trạng thái: Đang mượn, Đã trả",
        "Ghi nhận ngày trả và ghi chú khi trả",
        "Cảnh báo thiết bị đang được mượn khi có thao tác khác"
    ]
    
    for feature in loan_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    h2 = doc.add_heading('4.5. Quản lý đơn vị (Unit Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống hỗ trợ quản lý cấu trúc cây đơn vị theo cấp:").bold = True
    
    unit_features = [
        "Tạo cây đơn vị nhiều cấp (Cấp 0 là cao nhất)",
        "Quản lý thông tin: Tên, mã đơn vị, chỉ huy, địa chỉ, số điện thoại",
        "Thiết lập quan hệ cấp trên - cấp dưới",
        "Hiển thị dạng cây (TreeView) trực quan",
        "Gán thiết bị cho đơn vị cụ thể"
    ]
    
    for feature in unit_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    h2 = doc.add_heading('4.6. Quản lý người dùng (User Management)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Hệ thống phân quyền theo vai trò:").bold = True
    
    role_table = doc.add_table(rows=6, cols=3)
    role_table.style = 'Table Grid'
    
    role_data = [
        ('Vai trò', 'Mô tả', 'Quyền hạn'),
        ('Superadmin', 'Quản trị cao cấp', 'Toàn quyền hệ thống'),
        ('Admin', 'Quản trị viên', 'Quản lý user, đơn vị, danh mục'),
        ('Manager', 'Quản lý kho', 'CRUD thiết bị, xuất báo cáo'),
        ('Technician', 'Nhân viên kỹ thuật', 'Ghi nhật ký bảo dưỡng'),
        ('Viewer', 'Người xem', 'Chỉ xem, tra cứu'),
    ]
    
    for i, row_data in enumerate(role_data):
        cells = role_table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
            if i == 0:
                cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(cells[j], 'FFF3E0')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Các tính năng quản lý user:").italic = True
    
    user_features = [
        "Đăng nhập bắt buộc khi khởi động",
        "Mật khẩu được mã hóa SHA-256",
        "Đổi mật khẩu cá nhân",
        "Khóa/mở khóa tài khoản",
        "Tài khoản mặc định: admin/admin123"
    ]
    
    for feature in user_features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    h2 = doc.add_heading('4.7. Xuất báo cáo PDF (Report Export)', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    report_types = [
        ("Danh sách thiết bị", "Xuất toàn bộ hoặc theo bộ lọc, định dạng A4 ngang"),
        ("Bảng mã QR", "In nhiều mã QR trên một trang để cắt dán"),
        ("Hồ sơ thiết bị", "Chi tiết một thiết bị kèm lịch sử bảo dưỡng, cho mượn")
    ]
    
    for title, desc in report_types:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # ===== PHẦN 5: ĐIỂM NỔI BẬT =====
    h1 = doc.add_heading('5. ĐIỂM NỔI BẬT CỦA HỆ THỐNG', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    highlights = [
        ("🔒 Bảo mật cao - Hoạt động Offline", 
         "Phần mềm không yêu cầu kết nối Internet, dữ liệu được lưu trữ hoàn toàn cục bộ trong file SQLite. Điều này đảm bảo an toàn thông tin quân sự, tránh rò rỉ dữ liệu qua mạng."),
        
        ("📱 Ứng dụng công nghệ QR Code hiện đại",
         "Mỗi thiết bị được gán mã QR định danh duy nhất. Chỉ cần dùng webcam quét mã là có thể tra cứu ngay thông tin chi tiết, lịch sử bảo dưỡng, tình trạng cho mượn. Tiết kiệm thời gian tìm kiếm thủ công."),
        
        ("🎨 Giao diện thân thiện, hỗ trợ 2 chế độ Light/Dark",
         "Giao diện được thiết kế theo chuẩn UX hiện đại, dễ sử dụng. Người dùng có thể chuyển đổi giữa giao diện sáng và tối tùy theo sở thích và điều kiện ánh sáng."),
        
        ("🌳 Quản lý đơn vị theo cấu trúc cây",
         "Hệ thống cho phép xây dựng sơ đồ tổ chức đơn vị nhiều cấp, từ cấp cao nhất đến cấp thấp nhất. Việc gán thiết bị cho đơn vị trở nên chính xác và dễ quản lý."),
        
        ("👥 Phân quyền người dùng linh hoạt",
         "5 vai trò với quyền hạn khác nhau, từ người xem đến quản trị viên cao cấp. Đảm bảo mỗi người chỉ thao tác trong phạm vi được phép."),
        
        ("📊 Dashboard trực quan",
         "Trang tổng quan hiển thị các số liệu thống kê quan trọng: tổng thiết bị, phân loại theo tình trạng, theo danh mục, các hoạt động bảo dưỡng trong ngày."),
        
        ("📄 Xuất báo cáo PDF chuyên nghiệp",
         "Tự động tạo các báo cáo PDF với định dạng chuẩn, hỗ trợ font tiếng Việt, có thể in trực tiếp hoặc lưu file. Phục vụ tốt cho công tác báo cáo cấp trên."),
        
        ("💾 Dễ sao lưu và di chuyển",
         "Toàn bộ dữ liệu lưu trong một file database duy nhất, có thể sao chép sang máy tính khác dễ dàng. File PDF và ảnh QR được tổ chức gọn gàng trong thư mục data/."),
        
        ("🔧 Thiết kế mở rộng theo mô hình MVC",
         "Mã nguồn được tổ chức theo mô hình MVC chuẩn, dễ bảo trì và mở rộng thêm tính năng trong tương lai."),
        
        ("⚡ Hiệu năng cao, khởi động nhanh",
         "Phần mềm nhẹ, khởi động trong vài giây, không yêu cầu cấu hình máy tính cao. Hoạt động mượt mà ngay cả trên các máy tính cũ.")
    ]
    
    for title, desc in highlights:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        
        p = doc.add_paragraph(desc)
        p.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ===== PHẦN 6: HƯỚNG DẪN CÀI ĐẶT =====
    h1 = doc.add_heading('6. HƯỚNG DẪN CÀI ĐẶT VÀ SỬ DỤNG', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('6.1. Yêu cầu hệ thống', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    requirements = [
        "Hệ điều hành: Windows 10/11 (64-bit)",
        "Python: Phiên bản 3.10 trở lên",
        "RAM: Tối thiểu 4GB",
        "Ổ cứng: 100MB trống",
        "Webcam: Để sử dụng tính năng quét mã QR"
    ]
    
    for req in requirements:
        p = doc.add_paragraph(req, style='List Bullet')
    
    h2 = doc.add_heading('6.2. Các bước cài đặt', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    install_steps = [
        ("Bước 1: Tạo môi trường ảo", "python -m venv venv"),
        ("Bước 2: Kích hoạt môi trường", "venv\\Scripts\\activate"),
        ("Bước 3: Cài đặt thư viện", "pip install -r requirements.txt"),
        ("Bước 4: Chạy ứng dụng", "python main.py")
    ]
    
    for step, cmd in install_steps:
        p = doc.add_paragraph()
        run = p.add_run(step + ": ")
        run.bold = True
        p = doc.add_paragraph()
        run = p.add_run(cmd)
        run.font.name = 'Consolas'
        run.font.size = Pt(11)
        p.paragraph_format.left_indent = Cm(1)
    
    h2 = doc.add_heading('6.3. Hướng dẫn sử dụng cơ bản', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("Đăng nhập lần đầu:").bold = True
    p = doc.add_paragraph("Sử dụng tài khoản mặc định: admin / admin123")
    p.paragraph_format.left_indent = Cm(0.5)
    
    usage_guides = [
        ("Thêm thiết bị mới", "Vào mục 'Quản lý Trang bị' → Nhấn nút '+ Thêm mới' → Điền thông tin → Lưu"),
        ("Quét mã QR", "Vào mục 'Quét mã QR' → Chọn camera → Nhấn 'Bắt đầu quét' → Hướng camera về mã QR"),
        ("Ghi nhật ký bảo dưỡng", "Trong danh sách thiết bị, nhấn nút '📝' → Chọn loại công việc → Lưu"),
        ("Xuất báo cáo", "Vào menu 'Tệp' → Chọn loại báo cáo cần xuất")
    ]
    
    for title, guide in usage_guides:
        p = doc.add_paragraph()
        run = p.add_run(f"• {title}: ")
        run.bold = True
        p.add_run(guide)
    
    doc.add_page_break()
    
    # ===== PHẦN 7: KẾT LUẬN =====
    h1 = doc.add_heading('7. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', level=1)
    h1.runs[0].font.name = 'Times New Roman'
    h1.runs[0].font.size = Pt(16)
    
    h2 = doc.add_heading('7.1. Kết luận', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.add_run("""Phần mềm Quản lý Vũ khí Trang bị Kỹ thuật (QuanLyVKTBKT) đã được xây dựng hoàn chỉnh với đầy đủ các chức năng cần thiết cho công tác quản lý tại đơn vị:""")
    
    conclusions = [
        "Số hóa toàn bộ hồ sơ thiết bị, loại bỏ sổ sách giấy tờ",
        "Ứng dụng công nghệ QR Code giúp tra cứu nhanh chóng, chính xác",
        "Theo dõi đầy đủ lịch sử bảo dưỡng, cho mượn",
        "Hỗ trợ đắc lực cho công tác báo cáo, thống kê",
        "Đảm bảo an toàn thông tin với chế độ offline và phân quyền"
    ]
    
    for item in conclusions:
        p = doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run("""Phần mềm đã sẵn sàng để triển khai thử nghiệm tại các đơn vị, góp phần vào công cuộc chuyển đổi số trong quân đội.""").italic = True
    
    h2 = doc.add_heading('7.2. Hướng phát triển trong tương lai', level=2)
    h2.runs[0].font.name = 'Times New Roman'
    h2.runs[0].font.size = Pt(14)
    
    future_dev = [
        "Phát triển phiên bản mobile (Android/iOS) để quét QR bằng điện thoại",
        "Thêm tính năng đồng bộ dữ liệu giữa nhiều máy khi có mạng nội bộ",
        "Tích hợp nhận diện hư hỏng bằng AI từ ảnh chụp",
        "Thêm module quản lý kho vật tư, linh kiện thay thế",
        "Phát triển dashboard phân tích dữ liệu nâng cao",
        "Hỗ trợ in tem RFID thay vì QR cho môi trường khắc nghiệt"
    ]
    
    for item in future_dev:
        p = doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("─" * 40)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("© 2024 - Phần mềm Quản lý VKTBKT")
    run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Phiên bản 1.0.0")
    run.italic = True
    
    # Lưu file
    doc.save('docs/BaoCao_PhanMem_QuanLyVKTBKT.docx')
    print("✅ Đã tạo file báo cáo: docs/BaoCao_PhanMem_QuanLyVKTBKT.docx")

if __name__ == "__main__":
    # Tạo thư mục docs nếu chưa có
    os.makedirs('docs', exist_ok=True)
    create_report()